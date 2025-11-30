"""
Document Parser Service

Extracts text from various document formats:
- PDF (using pypdf)
- DOCX (using python-docx)
- TXT, MD (plain text)
- OCR support for scanned PDFs (local only)

Returns standardized Document objects with rich metadata.

OPTIMIZATIONS IMPLEMENTED:
1. Configurable PDF splitting (per-page vs combined)
2. OCR for image-only pages (local environment only)
3. Rich metadata (word count, upload date, file size, etc.)
4. Specific exception handling (corrupted PDF, permissions, etc.)
"""

from typing import List
from pathlib import Path
from datetime import datetime
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from docx import Document as DocxDocument
from langchain_core.documents import Document
import os

# OCR support (local only - Render has 512MB RAM limit)
try:
    if not os.getenv("RENDER"):
        import pytesseract
        from PIL import Image
        import pdf2image

        OCR_AVAILABLE = True
        print("[i] OCR support enabled (local environment)")
    else:
        OCR_AVAILABLE = False
except ImportError:
    OCR_AVAILABLE = False
    if not os.getenv("RENDER"):
        print("[!] OCR unavailable (install: pip install pytesseract pdf2image)")


class DocumentParser:
    """
    Parse documents and extract text content with rich metadata.

    Supports: PDF, DOCX, TXT, MD
    Returns: LangChain Document objects with metadata

    Features:
    - Configurable PDF splitting (per-page or combined)
    - OCR for scanned/image PDFs (local only)
    - Rich metadata (word count, upload date, file size)
    - Specific error handling
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """Check if file format is supported"""
        return Path(file_path).suffix.lower() in cls.SUPPORTED_EXTENSIONS

    @classmethod
    def parse(
        cls, file_path: str, split_by_page: bool = True, use_ocr: bool = False
    ) -> List[Document]:
        """
        Parse document and return LangChain Document objects.

        Args:
            file_path: Path to document file
            split_by_page: For PDFs, split into one doc per page (True) or combine (False)
            use_ocr: Use OCR for image-only pages (local only, requires tesseract)

        Returns:
            List of Document objects

        Raises:
            ValueError: If file format not supported
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if not cls.is_supported(file_path):
            raise ValueError(
                f"Unsupported file format: {path.suffix}. "
                f"Supported: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
            )

        # Route to appropriate parser
        extension = path.suffix.lower()

        if extension == ".pdf":
            return cls._parse_pdf(path, split_by_page, use_ocr)
        elif extension == ".docx":
            return cls._parse_docx(path)
        elif extension in {".txt", ".md"}:
            return cls._parse_text(path)

        raise ValueError(f"No parser for {extension}")

    @classmethod
    def _parse_pdf(
        cls, path: Path, split_by_page: bool = True, use_ocr: bool = False
    ) -> List[Document]:
        """
        Parse PDF file with flexible splitting and optional OCR.

        Args:
            split_by_page: True = one doc per page, False = combine all pages
            use_ocr: Try OCR on empty pages (local only)

        Returns:
            List of Document objects
        """
        try:
            reader = PdfReader(str(path))
            file_stat = path.stat()
            upload_date = datetime.fromtimestamp(file_stat.st_mtime).isoformat()

            if split_by_page:
                # Current behavior: One document per page
                documents = []

                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()

                    # Try OCR if page is empty and OCR available
                    if not text.strip() and use_ocr and OCR_AVAILABLE:
                        text = cls._ocr_pdf_page(path, page_num)
                        if text:
                            print(f"[i] OCR extracted text from page {page_num + 1}")

                    # Skip if still empty
                    if not text.strip():
                        continue

                    # Create Document with RICH metadata
                    doc = Document(
                        page_content=text,
                        metadata={
                            "source": str(path),
                            "file_name": path.name,
                            "file_type": "pdf",
                            "page": page_num + 1,
                            "total_pages": len(reader.pages),
                            "char_count": len(text),
                            "word_count": len(text.split()),
                            "upload_date": upload_date,
                            "file_size_kb": round(file_stat.st_size / 1024, 2),
                        },
                    )
                    documents.append(doc)

                return documents

            else:
                # NEW: Combine all pages into one document
                all_text = []
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()

                    # Try OCR if empty
                    if not text.strip() and use_ocr and OCR_AVAILABLE:
                        text = cls._ocr_pdf_page(path, page_num)

                    if text.strip():
                        all_text.append(text)

                combined_text = "\n\n".join(all_text)

                return [
                    Document(
                        page_content=combined_text,
                        metadata={
                            "source": str(path),
                            "file_name": path.name,
                            "file_type": "pdf",
                            "total_pages": len(reader.pages),
                            "char_count": len(combined_text),
                            "word_count": len(combined_text.split()),
                            "upload_date": upload_date,
                            "file_size_kb": round(file_stat.st_size / 1024, 2),
                            "combined": True,  # Flag that pages were combined
                        },
                    )
                ]

        except PdfReadError as e:
            raise ValueError(f"PDF corrupted or encrypted: {path.name} - {e}")
        except PermissionError:
            raise ValueError(f"Permission denied: {path.name}")
        except Exception as e:
            raise ValueError(f"Failed to parse PDF {path.name}: {e}")

    @classmethod
    def _ocr_pdf_page(cls, path: Path, page_num: int) -> str:
        """
        OCR a single PDF page (local only).

        Returns extracted text or empty string if OCR fails.
        """
        if not OCR_AVAILABLE:
            return ""

        try:
            # Convert PDF page to image
            images = pdf2image.convert_from_path(
                str(path), first_page=page_num + 1, last_page=page_num + 1
            )

            if not images:
                return ""

            # OCR the image
            text = pytesseract.image_to_string(images[0])
            return text.strip()

        except Exception as e:
            print(f"[!] OCR failed for page {page_num + 1}: {e}")
            return ""

    @classmethod
    def _parse_docx(cls, path: Path) -> List[Document]:
        """
        Parse DOCX file.

        Returns single Document with all content and rich metadata.
        """
        try:
            doc = DocxDocument(str(path))

            # Extract all paragraphs
            text = "\n".join([para.text for para in doc.paragraphs])

            if not text.strip():
                raise ValueError(f"DOCX file is empty: {path.name}")

            # Rich metadata
            file_stat = path.stat()

            return [
                Document(
                    page_content=text,
                    metadata={
                        "source": str(path),
                        "file_name": path.name,
                        "file_type": "docx",
                        "paragraphs": len(doc.paragraphs),
                        "char_count": len(text),
                        "word_count": len(text.split()),
                        "upload_date": datetime.fromtimestamp(
                            file_stat.st_mtime
                        ).isoformat(),
                        "file_size_kb": round(file_stat.st_size / 1024, 2),
                    },
                )
            ]

        except Exception as e:
            raise ValueError(f"Failed to parse DOCX {path.name}: {e}")

    @classmethod
    def _parse_text(cls, path: Path) -> List[Document]:
        """
        Parse plain text file (TXT, MD).

        Returns single Document with all content and rich metadata.
        """
        try:
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()

            if not text.strip():
                raise ValueError(f"Text file is empty: {path.name}")

            # Rich metadata
            file_stat = path.stat()

            return [
                Document(
                    page_content=text,
                    metadata={
                        "source": str(path),
                        "file_name": path.name,
                        "file_type": path.suffix[1:],  # Remove the dot
                        "char_count": len(text),
                        "word_count": len(text.split()),
                        "upload_date": datetime.fromtimestamp(
                            file_stat.st_mtime
                        ).isoformat(),
                        "file_size_kb": round(file_stat.st_size / 1024, 2),
                    },
                )
            ]

        except UnicodeDecodeError:
            # Try different encoding
            try:
                with open(path, "r", encoding="latin-1") as f:
                    text = f.read()

                file_stat = path.stat()

                return [
                    Document(
                        page_content=text,
                        metadata={
                            "source": str(path),
                            "file_name": path.name,
                            "file_type": path.suffix[1:],
                            "encoding": "latin-1",
                            "char_count": len(text),
                            "word_count": len(text.split()),
                            "upload_date": datetime.fromtimestamp(
                                file_stat.st_mtime
                            ).isoformat(),
                            "file_size_kb": round(file_stat.st_size / 1024, 2),
                        },
                    )
                ]
            except Exception as e:
                raise ValueError(f"Failed to parse text file {path.name}: {e}")

        except Exception as e:
            raise ValueError(f"Failed to parse text file {path.name}: {e}")


# =============================================================================
# Test Parser
# =============================================================================
if __name__ == "__main__":
    print("=" * 70)
    print("Document Parser Test (With Optimizations)")
    print("=" * 70)

    # Test with sample text file
    test_content = """
    RAG (Retrieval-Augmented Generation) is a technique for enhancing LLM responses.

    It works by:
    1. Retrieving relevant documents from a knowledge base
    2. Passing those documents as context to the LLM
    3. Generating an answer based on the retrieved context

    This improves accuracy and allows LLMs to access specific knowledge.
    """

    # Create sample file
    test_file = Path("../data/test_document.txt")
    test_file.parent.mkdir(exist_ok=True)
    test_file.write_text(test_content)

    # Test parser
    print(f"\nParsing: {test_file.name}")
    print("-" * 70)

    documents = DocumentParser.parse(str(test_file))

    for doc in documents:
        print(f"Content: {doc.page_content[:100]}...")
        print(f"Metadata: {doc.metadata}")
        print()

    # Cleanup
    test_file.unlink()

    print("=" * 70)
    print("[OK] Parser working with optimizations!")
    print("=" * 70)
    print("\nOptimizations enabled:")
    print("  ✅ Configurable PDF splitting")
    print("  ✅ Rich metadata (word count, file size, upload date)")
    print("  ✅ Specific error handling")
    print(f"  {'✅' if OCR_AVAILABLE else '⏸️'} OCR support (local only)")
